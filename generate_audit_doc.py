from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin   = Cm(2.2)
section.right_margin  = Cm(2.2)

# ── Helper: set paragraph shading ─────────────────────────────────────────────
def shade_paragraph(para, hex_color):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    pPr.append(shd)

def shade_cell(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'bottom', 'start', 'end', 'left', 'right'):
        tag = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'),   kwargs.get('val',   'single'))
        tag.set(qn('w:sz'),    kwargs.get('sz',    '4'))
        tag.set(qn('w:space'), '0')
        tag.set(qn('w:color'), kwargs.get('color', '000000'))
        tcBorders.append(tag)
    tcPr.append(tcBorders)

# ── TITLE PAGE ────────────────────────────────────────────────────────────────
title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
shade_paragraph(title_para, '1A237E')
run = title_para.add_run('UDAAN Admin Panel')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

sub_para = doc.add_paragraph()
sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
shade_paragraph(sub_para, '1A237E')
run2 = sub_para.add_run('Deep Scan Audit Report — 25 Issues Found')
run2.bold = True
run2.font.size = Pt(16)
run2.font.color.rgb = RGBColor(0xFF, 0xCC, 0x00)

date_para = doc.add_paragraph()
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
shade_paragraph(date_para, '283593')
date_run = date_para.add_run(f'  Scanned: {datetime.date.today().strftime("%B %d, %Y")}  |  Static Code Analysis + Live Server Review  ')
date_run.font.size = Pt(11)
date_run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xFF)

doc.add_paragraph()

# ── LEGEND ────────────────────────────────────────────────────────────────────
legend_heading = doc.add_paragraph()
run = legend_heading.add_run('Severity Legend')
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

legend_table = doc.add_table(rows=1, cols=4)
legend_table.alignment = WD_TABLE_ALIGNMENT.LEFT
legend_table.style = 'Table Grid'
hdr = legend_table.rows[0].cells
labels = [
    ('🔴  CRITICAL', 'FF1744', 'FFEBEE'),
    ('🟠  HIGH',     'FF6D00', 'FFF3E0'),
    ('🟡  MEDIUM',   'F9A825', 'FFFDE7'),
    ('🔵  LOW',      '1565C0', 'E3F2FD'),
]
for i, (text, border, bg) in enumerate(labels):
    cell = hdr[i]
    shade_cell(cell, bg)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(
        int(border[0:2], 16), int(border[2:4], 16), int(border[4:6], 16)
    )

doc.add_paragraph()

# ── ISSUES DATA ───────────────────────────────────────────────────────────────
issues = [
    # (num, severity_label, severity_bg, severity_txt, title, file_ref, description)
    (
        1, '🔴 CRITICAL', 'FFEBEE', 'FF1744',
        'Hardcoded SECRET_KEY in settings.py',
        'project/settings.py — Line 29',
        (
            'The Django SECRET_KEY is hardcoded directly in source code with the insecure prefix '
            '"django-insecure-...". Anyone with repository access can forge session tokens, '
            'CSRF tokens, and signed cookies — a complete authentication bypass risk.\n\n'
            'Code:\n'
            '    SECRET_KEY = "django-insecure-j0rm1^nre3hpqku&6%uu..."\n\n'
            'Fix: Move to a .env file and load via os.environ.get("SECRET_KEY"). '
            'Rotate the key immediately after moving it.'
        ),
    ),
    (
        2, '🔴 CRITICAL', 'FFEBEE', 'FF1744',
        'Unauthenticated Donor PII Exposure via Search API',
        'blood_request/views.py — Lines 76–99',
        (
            'The search_donors endpoint has NO @login_required decorator and NO rate limiting. '
            'It returns full PII — name, phone number, email, city and state — for every matching '
            'donor to any unauthenticated visitor.\n\n'
            'Impact: An attacker can enumerate the entire donor database by looping through blood '
            'groups or cities. This is a GDPR/DPDP Act violation.\n\n'
            'Fix: Add @login_required and @ratelimit(key="ip", rate="20/h"). '
            'Consider masking phone/email for non-manager staff.'
        ),
    ),
    (
        3, '🔴 CRITICAL', 'FFEBEE', 'FF1744',
        'BloodRequest Status is Permanently Stuck — FSM Transitions Missing in Admin',
        'blood_request/admin.py — Lines 88–92',
        (
            'The BloodRequest model uses a FSMField(protected=True) with states: '
            'Received → Verified → Fulfilling → Closed. '
            'In the admin panel the status field is set as readonly_fields = ("status",), '
            'but there are NO custom admin actions to trigger FSM transitions.\n\n'
            'Impact: Staff cannot change any blood request status from the admin panel. '
            'All requests stay forever in "Received" state.\n\n'
            'Fix: Add custom ModelAdmin actions:\n'
            '    def mark_verified(self, request, queryset): ...\n'
            '    def mark_fulfilling(self, request, queryset): ...\n'
            '    def mark_closed(self, request, queryset): ...'
        ),
    ),
    (
        4, '🟠 HIGH', 'FFF3E0', 'FF6D00',
        'NameError Crash in add_task_comment — messages Not Imported at Usage Point',
        'blood_request/views.py — Lines 339, 341',
        (
            'The add_task_comment view uses "messages.success()" and "messages.error()" '
            'at lines 339 and 341, but the import "from django.contrib import messages" '
            'only appears at line 820 (later in the file). In Python, function bodies are '
            'executed at call time, but the name "messages" will resolve to whichever value '
            'was in scope when the module loaded. Because the import is below the function, '
            'this raises a NameError at runtime when a comment is posted.\n\n'
            'Fix: Add "from django.contrib import messages" to the top-level imports at the '
            'top of views.py.'
        ),
    ),
    (
        5, '🟠 HIGH', 'FFF3E0', 'FF6D00',
        '10+ Models Missing from Admin Panel — Completely Invisible to Staff',
        'blood_request/admin.py',
        (
            'The following models exist in models.py and are actively used in views, '
            'but are NOT registered in admin.py. Staff cannot manage them via the admin panel:\n\n'
            '  • Donation — Cannot record or view blood donations\n'
            '  • SubTask — Cannot manage sub-tasks\n'
            '  • TaskComment — Cannot moderate comments\n'
            '  • Team — Cannot manage teams\n'
            '  • SharedNote — Cannot manage wiki notes\n'
            '  • Workspace / WorkspaceMember — Cannot manage workspaces\n'
            '  • Expense — Cannot manage expense records\n'
            '  • TaskAutomationRule — Cannot configure automation rules\n'
            '  • NewsletterSubscription — Cannot see subscribers\n'
            '  • CampusAmbassadorApplication — Public applications are invisible\n\n'
            'Fix: Register each model using @admin.register(ModelName) in admin.py.'
        ),
    ),
    (
        6, '🟠 HIGH', 'FFF3E0', 'FF6D00',
        'Duplicate our_policies View Function Definition',
        'blood_request/views.py — Lines 1187 & 1190',
        (
            'The our_policies function is defined TWICE in views.py. The first definition '
            '(line 1187) renders the template with no context data (empty page). The second '
            '(line 1190) correctly fetches PolicyReport objects and passes them to the template.\n\n'
            'Python silently uses the second definition, but the first is silent dead code that '
            'will cause confusion and may be accidentally activated if someone deletes the second.\n\n'
            'Fix: Delete the first (line 1187) empty definition. Keep only the data-fetching one.'
        ),
    ),
    (
        7, '🟠 HIGH', 'FFF3E0', 'FF6D00',
        'User Creation in Portal Bypasses Django Password Validators',
        'blood_request/views.py — Lines 1075–1103',
        (
            'The user_add portal view creates users via:\n'
            '    user = User.objects.create_user(username=username, email=email, password=password)\n\n'
            'This bypasses all AUTH_PASSWORD_VALIDATORS configured in settings.py '
            '(similarity check, minimum length, common password, numeric-only checks). '
            'There is also no password confirmation field, so typos create broken accounts.\n\n'
            'Fix: Use Django\'s UserCreationForm or manually call '
            'validate_password(password, user) before creating the user. '
            'Add a confirm_password field to the HTML form.'
        ),
    ),
    (
        8, '🟠 HIGH', 'FFF3E0', 'FF6D00',
        'TIME_ZONE Set to UTC Instead of IST (Asia/Kolkata)',
        'project/settings.py — Line 190',
        (
            'TIME_ZONE = "UTC" means all timestamps stored and displayed in the admin panel '
            'are in UTC (+00:00) rather than IST (+05:30). For an Indian NGO this means:\n\n'
            '  • Task due dates shown 5 hours 30 minutes early\n'
            '  • Appointment times off by 5.5 hours\n'
            '  • BloodRequest created_at times off by 5.5 hours\n'
            '  • Interaction logs show wrong time\n'
            '  • Date hierarchy filters in admin are incorrect\n\n'
            'Fix: Change to TIME_ZONE = "Asia/Kolkata"'
        ),
    ),
    (
        9, '🟠 HIGH', 'FFF3E0', 'FF6D00',
        'InteractionAdmin list_filter on ForeignKey staff — Loads ALL Users Into Sidebar',
        'blood_request/admin.py — Line 73',
        (
            'list_filter = ("staff", "interaction_type", "outcome") causes Django admin to '
            'load all User objects into a sidebar list filter. As user count grows this '
            'renders slowly and becomes unusable.\n\n'
            'Fix: Remove "staff" from list_filter and instead add a search_fields entry:\n'
            '    search_fields = ("staff__username", "notes")'
        ),
    ),
    (
        10, '🟠 HIGH', 'FFF3E0', 'FF6D00',
        'CampusAmbassadorApplication Imported But Never Registered in Admin',
        'blood_request/admin.py — Lines 111–116',
        (
            'In admin.py, the CampusAmbassadorApplication model is imported but never '
            'registered with @admin.register:\n\n'
            '    from .models import CampusAmbassador, CampusAmbassadorApplication\n'
            '    # Only CampusAmbassador is registered below!\n\n'
            'Impact: Every application submitted via the public Campus Ambassador form is '
            'invisible in the admin panel. Staff cannot review, accept, or reject applications.\n\n'
            'Fix: Add:\n'
            '    @admin.register(CampusAmbassadorApplication)\n'
            '    class CampusAmbassadorApplicationAdmin(admin.ModelAdmin):\n'
            '        list_display = ("full_name","college","city","status","applied_at")\n'
            '        list_filter = ("status",)\n'
            '        list_editable = ("status",)'
        ),
    ),
    (
        11, '🟡 MEDIUM', 'FFFDE7', 'F9A825',
        '10+ Views Defined Without URL Routes — Pages Are Unreachable',
        'blood_request/urls.py + views.py',
        (
            'The following views are fully implemented but have no URL routes defined in '
            'blood_request/urls.py, making them unreachable by users:\n\n'
            '  • campaign_list\n'
            '  • project_list\n'
            '  • project_detail\n'
            '  • donate_page\n'
            '  • blogs_page\n'
            '  • news_clippings\n'
            '  • workplace_living\n'
            '  • volunteering\n'
            '  • resources_page\n'
            '  • manager_dashboard\n\n'
            'Fix: Add path(...) entries in urls.py for each missing view, '
            'or verify they exist in project/urls.py.'
        ),
    ),
    (
        12, '🟡 MEDIUM', 'FFFDE7', 'F9A825',
        'TaskForm Missing project and dependencies Fields',
        'blood_request/forms.py — Lines 31–40',
        (
            'The TaskForm used in the portal task-create view only has:\n'
            '    fields = ["title","description","priority","status","assigned_to","due_date","recurrence_rule"]\n\n'
            'But TaskCreateView and TaskUpdateView (views.py L1242) include:\n'
            '    fields = ["title","project","description","assigned_to","status","priority",'
            '"due_date","recurrence_rule","dependencies"]\n\n'
            'Impact: When managers create tasks via the portal form, they cannot link them '
            'to a project or set task dependencies.\n\n'
            'Fix: Add project and dependencies to TaskForm fields list.'
        ),
    ),
    (
        13, '🟡 MEDIUM', 'FFFDE7', 'F9A825',
        'Donation Model Not Registered in Admin',
        'blood_request/admin.py',
        (
            'The Donation model (tracks individual blood donations per donor) is defined in '
            'models.py and is displayed in the donor_detail portal view. However it is '
            'completely absent from admin.py.\n\n'
            'Impact: Staff cannot record, view, edit, or delete donations from the admin panel. '
            'The only way to manage donations is through the custom portal donor detail view.\n\n'
            'Fix:\n'
            '    @admin.register(Donation)\n'
            '    class DonationAdmin(admin.ModelAdmin):\n'
            '        list_display = ("donor","date","units","created_at")\n'
            '        search_fields = ("donor__name",)\n'
            '        list_filter = ("date",)'
        ),
    ),
    (
        14, '🟡 MEDIUM', 'FFFDE7', 'F9A825',
        'BlogAdmin Missing Search Fields, Filters, and Rich Text Editor',
        'blood_request/admin.py — Lines 94–97',
        (
            'The BlogAdmin class only has:\n'
            '    list_display = ("title", "created_at")\n\n'
            'Problems:\n'
            '  1. No search_fields — cannot search blogs by title or content\n'
            '  2. No list_filter — cannot filter by date\n'
            '  3. Blog.content uses CKEditor5Field in the model, but BlogAdmin does not '
            '     configure the widget — the admin shows a plain <textarea> instead of '
            '     the rich text editor\n\n'
            'Fix: Add search_fields, list_filter, and configure formfield_overrides to '
            'use CKEditor5Widget for the content field.'
        ),
    ),
    (
        15, '🟡 MEDIUM', 'FFFDE7', 'F9A825',
        'ProjectAdmin Missing Key Fields in list_display',
        'blood_request/admin.py — Lines 83–86',
        (
            'ProjectAdmin only shows:\n'
            '    list_display = ("title", "created_at")\n\n'
            'Missing from the list view:\n'
            '  • date (when the project happened)\n'
            '  • managers (who owns it)\n'
            '  • slug (for URL generation)\n\n'
            'Staff have no way to distinguish between projects or find their managers '
            'without clicking into each one individually.\n\n'
            'Fix:\n'
            '    list_display = ("title", "date", "created_at")\n'
            '    search_fields = ("title", "description")\n'
            '    list_filter = ("date",)\n'
            '    filter_horizontal = ("managers",)'
        ),
    ),
    (
        16, '🟡 MEDIUM', 'FFFDE7', 'F9A825',
        'logged_by Field Exposed in ExpenseCreateView — Allows Spoofing',
        'blood_request/views.py — Line 1320',
        (
            'The ExpenseCreateView includes logged_by in its fields list:\n'
            '    fields = ["title","amount","date","category","campaign","project","logged_by",...]\n\n'
            'This means any manager can create an expense record and attribute it to any other '
            'user, breaking the audit trail.\n\n'
            'Fix: Remove logged_by from fields and auto-set it in form_valid:\n'
            '    def form_valid(self, form):\n'
            '        form.instance.logged_by = self.request.user\n'
            '        return super().form_valid(form)'
        ),
    ),
    (
        17, '🟡 MEDIUM', 'FFFDE7', 'F9A825',
        'ContactMessage Admin Has No Reply or Delete-After-Read Action',
        'blood_request/admin.py — Lines 124–133',
        (
            'The ContactMessageAdmin only has a "mark as read" action. Staff have no way '
            'to:\n'
            '  • Reply to the sender from the admin panel\n'
            '  • Bulk delete read messages (inbox cleanup)\n'
            '  • See the full message body in the list view\n\n'
            'Additionally, there is no message preview in list_display — staff must click '
            'each message to read it.\n\n'
            'Fix: Add list_display = ("first_name","email","subject","message_preview","is_read","created_at") '
            'with a short message_preview method, and add a mark_as_unread and delete_read_messages action.'
        ),
    ),
    (
        18, '🟡 MEDIUM', 'FFFDE7', 'F9A825',
        'send_digest_portal View: Uses settings but Missing Import',
        'blood_request/views.py — Line 1477',
        (
            'In send_digest_portal (line 1477), the code references:\n'
            '    from_email=settings.DEFAULT_FROM_EMAIL\n\n'
            'However, "settings" is not imported at the top of views.py. '
            'The Django settings object is referenced by name "settings" which is only '
            'available if "from django.conf import settings" is present.\n\n'
            'This will cause a NameError when a manager tries to send the daily digest.\n\n'
            'Fix: Add "from django.conf import settings" to the top of views.py.'
        ),
    ),
    (
        19, '🟡 MEDIUM', 'FFFDE7', 'F9A825',
        'Appointment admin_url in calendar_events_api Points to Wrong Route',
        'blood_request/views.py — Line 797',
        (
            'In calendar_events_api, task events link to:\n'
            '    "url": "/admin/portal/"  # linking to dashboard for now\n\n'
            'This is a placeholder comment that was never updated. Clicking a task event '
            'in the FullCalendar widget takes users to the generic portal dashboard '
            'instead of the specific task detail page.\n\n'
            'Fix: Change to:\n'
            '    "url": f"/admin/portal/task/{task.id}/"'
        ),
    ),
    (
        20, '🟡 MEDIUM', 'FFFDE7', 'F9A825',
        'Two Conflicting Path("") Routes in urls.py — index and home_view Both on Root',
        'blood_request/urls.py — Lines 4 & 42',
        (
            'In urls.py there are two routes both mapped to path(""):\n'
            '    path("", views.index, name="index")    # Line 4\n'
            '    path("", views.home_view, name="home") # Line 42\n\n'
            'Django uses the first match, so home_view (the actual homepage with dynamic '
            'content) is never reached — the bare index view (which just returns a static '
            'template) always wins.\n\n'
            'Fix: Remove the duplicate path("", views.index) or rename it to a different path '
            'if index.html is still needed (e.g. path("old-index/", views.index)).'
        ),
    ),
    (
        21, '🔵 LOW', 'E3F2FD', '1565C0',
        'BloodDonor score and donation_count Are readonly — No Way to Correct Corruption',
        'blood_request/admin.py — Line 81',
        (
            'BloodDonorAdmin sets:\n'
            '    readonly_fields = ("score", "donation_count")\n\n'
            'These are calculated fields, but if a donor\'s score or count becomes corrupted '
            '(e.g., via a manual DB edit, migration error, or bug), there is no admin action '
            'to recalculate or override them.\n\n'
            'Fix: Add a recalculate_score custom admin action that recomputes the score from '
            'the related Donation records.'
        ),
    ),
    (
        22, '🔵 LOW', 'E3F2FD', '1565C0',
        'JobPosting list_editable Without Confirmation — Risk of Accidental Mass Deactivation',
        'blood_request/admin.py — Line 151',
        (
            'JobPostingAdmin uses:\n'
            '    list_editable = ("is_active",)\n\n'
            'This renders is_active as a checkbox on every row of the job list. '
            'A manager can accidentally uncheck multiple jobs and submit, deactivating them '
            'all without any confirmation prompt.\n\n'
            'Fix: Remove is_active from list_editable and instead provide a toggle action '
            'with confirmation:\n'
            '    actions = ["activate_jobs", "deactivate_jobs"]'
        ),
    ),
    (
        23, '🔵 LOW', 'E3F2FD', '1565C0',
        'InteractionAdmin search_fields Only Searches notes Field',
        'blood_request/admin.py — Line 74',
        (
            'search_fields = ("notes",)\n\n'
            'This means staff can only search interactions by notes text. '
            'They cannot search by the staff member\'s username, interaction type, or outcome.\n\n'
            'Fix:\n'
            '    search_fields = ("notes", "staff__username", "staff__first_name")'
        ),
    ),
    (
        24, '🔵 LOW', 'E3F2FD', '1565C0',
        'AnnouncementAdmin Missing Priority / Expiry Date Fields',
        'blood_request/admin.py — Lines 63–66',
        (
            'AnnouncementAdmin shows:\n'
            '    list_display = ("title", "is_active", "created_at")\n\n'
            'The Announcement model has no expiry_date or priority field, which means:\n'
            '  • Old announcements accumulate forever (no auto-expiry)\n'
            '  • All announcements have equal weight regardless of urgency\n\n'
            'Recommendation: Add expiry_date = models.DateField(null=True, blank=True) '
            'and priority field to the Announcement model, then update AnnouncementAdmin '
            'and the staff_dashboard view to filter out expired announcements.'
        ),
    ),
    (
        25, '🔵 LOW', 'E3F2FD', '1565C0',
        'No admin.site.site_header / site_title Customization — Shows Generic "Django Administration"',
        'blood_request/admin.py or project/urls.py',
        (
            'The Django admin panel displays the generic "Django administration" header and '
            'tab title. There is no branding customization for UDAAN.\n\n'
            'Fix: Add to admin.py or urls.py:\n'
            '    from django.contrib import admin\n'
            '    admin.site.site_header = "UDAAN Administration"\n'
            '    admin.site.site_title  = "UDAAN Admin Portal"\n'
            '    admin.site.index_title = "Welcome to UDAAN Admin Dashboard"'
        ),
    ),
]

# ── SUMMARY TABLE ─────────────────────────────────────────────────────────────
heading = doc.add_paragraph()
r = heading.add_run('Executive Summary')
r.bold = True
r.font.size = Pt(16)
r.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

counts = {'🔴 CRITICAL': 0, '🟠 HIGH': 0, '🟡 MEDIUM': 0, '🔵 LOW': 0}
for issue in issues:
    counts[issue[1]] += 1

summary_table = doc.add_table(rows=2, cols=4)
summary_table.style = 'Table Grid'
summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
colors = [
    ('🔴 CRITICAL', 'FF1744', 'FFEBEE'),
    ('🟠 HIGH',     'FF6D00', 'FFF3E0'),
    ('🟡 MEDIUM',   'F9A825', 'FFFDE7'),
    ('🔵 LOW',      '1565C0', 'E3F2FD'),
]
for col_i, (sev, txt_hex, bg_hex) in enumerate(colors):
    hdr_cell = summary_table.rows[0].cells[col_i]
    data_cell = summary_table.rows[1].cells[col_i]

    shade_cell(hdr_cell, bg_hex)
    p = hdr_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(sev)
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(
        int(txt_hex[0:2],16), int(txt_hex[2:4],16), int(txt_hex[4:6],16)
    )

    shade_cell(data_cell, bg_hex)
    p2 = data_cell.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(str(counts[sev]) + ' issue(s)')
    r2.bold = True
    r2.font.size = Pt(14)
    r2.font.color.rgb = RGBColor(
        int(txt_hex[0:2],16), int(txt_hex[2:4],16), int(txt_hex[4:6],16)
    )

doc.add_paragraph()

# ── ISSUE SECTIONS ─────────────────────────────────────────────────────────────
sev_colors = {
    '🔴 CRITICAL': ('FF1744', 'FFEBEE', '1A237E'),
    '🟠 HIGH':     ('FF6D00', 'FFF3E0', '1A237E'),
    '🟡 MEDIUM':   ('F9A825', 'FFFDE7', '1A237E'),
    '🔵 LOW':      ('1565C0', 'E3F2FD', '1A237E'),
}

current_sev = None
for issue in issues:
    num, sev, sev_bg, sev_txt_hex, title, file_ref, description = issue

    # Section header when severity changes
    if sev != current_sev:
        current_sev = sev
        sec_para = doc.add_paragraph()
        shade_paragraph(sec_para, sev_colors[sev][0])
        r = sec_para.add_run(f'  {sev} ISSUES  ')
        r.bold = True
        r.font.size = Pt(13)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        doc.add_paragraph()

    # Issue title bar
    title_p = doc.add_paragraph()
    shade_paragraph(title_p, 'E8EAF6')
    r1 = title_p.add_run(f'  Issue #{num}  ')
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    r1_bg = title_p.add_run()
    # badge color in the run
    badge = title_p.add_run(f' {sev} ')
    badge.bold = True
    badge.font.size = Pt(10)
    badge.font.color.rgb = RGBColor(
        int(sev_colors[sev][0][0:2],16),
        int(sev_colors[sev][0][2:4],16),
        int(sev_colors[sev][0][4:6],16),
    )
    title_r = title_p.add_run(f'  {title}')
    title_r.bold = True
    title_r.font.size = Pt(12)
    title_r.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

    # File reference
    file_p = doc.add_paragraph()
    file_p.paragraph_format.left_indent = Inches(0.2)
    rf = file_p.add_run(f'📁  File: {file_ref}')
    rf.italic = True
    rf.font.size = Pt(10)
    rf.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # Description
    desc_p = doc.add_paragraph()
    desc_p.paragraph_format.left_indent = Inches(0.2)
    rd = desc_p.add_run(description)
    rd.font.size = Pt(10.5)

    doc.add_paragraph()

# ── FOOTER NOTE ───────────────────────────────────────────────────────────────
footer_para = doc.add_paragraph()
shade_paragraph(footer_para, '1A237E')
rf = footer_para.add_run(
    '  UDAAN Admin Panel Audit | Generated by Antigravity AI | '
    f'{datetime.date.today().strftime("%B %d, %Y")}  '
)
rf.bold = True
rf.font.size = Pt(9)
rf.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

# ── SAVE ──────────────────────────────────────────────────────────────────────
output_path = r'D:\UDAAN-main\UDAAN_Admin_Audit_Report.docx'
doc.save(output_path)
print(f'[OK] Document saved: {output_path}')
